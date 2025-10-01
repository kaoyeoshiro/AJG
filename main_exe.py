# main_exe.py — v3 completo (Versão para executável)
# -*- coding: utf-8 -*-
"""
App Tkinter para:
1) Consultar XML do TJ-MS (SOAP) com credenciais do config.py;
2) Extrair partes (AT/PA + AJG) e decisões/ despachos (3/11009) com complementos;
3) Detectar classe de cumprimento e indício de apenso/apensado (aviso apenas);
4) Enviar JSON para LLM via OpenRouter e gerar relatório;
5) UI com log embutido, teste integral, visualização do JSON e fallback defensivo.

Requisitos: requests (Tkinter vem no Python oficial Win/Mac; no Linux instalar via sistema)
"""

import os
import sys
import re
import json
import html
import logging
import threading
import base64
from datetime import datetime
from typing import Tuple, List, Dict, Any

import requests
from requests.adapters import HTTPAdapter, Retry
import xml.etree.ElementTree as ET

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText
from tkinter import font as tkfont

# =========================
# Importa configurações do config.py
# =========================
from config import (
    TJ_WSDL_URL, TJ_WS_USER, TJ_WS_PASS,
    OPENROUTER_API_KEY, OPENROUTER_ENDPOINT, DEFAULT_MODEL,
    STRICT_CNJ_CHECK, CLASSES_CUMPRIMENTO, NS
)

# =========================
# Importa módulo de auto-atualização
# =========================
try:
    from scripts.updater import check_and_update
    UPDATER_AVAILABLE = True
except ImportError:
    try:
        # Fallback para importação antiga (compatibilidade)
        from updater import check_and_update  # type: ignore[import-not-found]
        UPDATER_AVAILABLE = True
    except ImportError:
        UPDATER_AVAILABLE = False
        print("Modulo updater nao encontrado - funcionalidade de auto-update desabilitada")

# =========================
# Importa gerenciador de chaves
# =========================
try:
    from scripts.key_manager import get_api_key, KeyManager
    KEY_MANAGER_AVAILABLE = True
except ImportError:
    try:
        # Fallback para importação antiga (compatibilidade)
        from key_manager import get_api_key, KeyManager  # type: ignore[import-not-found]
        KEY_MANAGER_AVAILABLE = True
    except ImportError:
        KEY_MANAGER_AVAILABLE = False
        print("Modulo key_manager nao encontrado - usando configuracao estatica")

# =========================
# Logging (terminal)
# =========================
logger = logging.getLogger("RelatorioTJMS")
logger.setLevel(logging.INFO)
_ch = logging.StreamHandler()
_ch.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(message)s"))
logger.addHandler(_ch)

# =========================
# Utilitários gerais
# =========================
def make_session() -> requests.Session:
    s = requests.Session()
    retry = Retry(total=4, backoff_factor=0.6, status_forcelist=[429, 500, 502, 503, 504])
    s.mount("http://", HTTPAdapter(max_retries=retry))
    s.mount("https://", HTTPAdapter(max_retries=retry))

    # Configurar certificados SSL para executável PyInstaller
    if getattr(sys, 'frozen', False):
        # Executável empacotado - usar cacert.pem incluído
        import os
        bundle_path = sys._MEIPASS
        cacert_path = os.path.join(bundle_path, 'certifi', 'cacert.pem')
        if os.path.exists(cacert_path):
            s.verify = cacert_path
            logger.info(f"Usando certificados SSL de: {cacert_path}")
        else:
            # Fallback: usar certifi padrão
            import certifi
            s.verify = certifi.where()
            logger.info(f"Fallback: usando certificados SSL de: {certifi.where()}")

    return s

def only_digits(s: str) -> str:
    return re.sub(r"\D", "", s or "")

def format_cnj(num: str) -> str:
    d = only_digits(num)
    if len(d) != 20:
        return num
    return f"{d[0:7]}-{d[7:9]}.{d[9:13]}.{d[13:14]}.{d[14:16]}.{d[16:20]}"

def cnj_checksum_ok(d: str) -> bool:
    # Checagem simplificada (se quiser, troque pelo algoritmo 97-10 completo)
    if len(d) != 20 or not d.isdigit():
        return False
    ano = int(d[9:13])
    return 1900 <= ano <= datetime.now().year + 1

def validate_config() -> Tuple[bool, str]:
    miss = []
    if not TJ_WSDL_URL: miss.append("TJ_WSDL_URL")
    if not TJ_WS_USER:  miss.append("TJ_WS_USER")
    if not TJ_WS_PASS:  miss.append("TJ_WS_PASS")
    # Permite placeholder durante build/desenvolvimento
    if not OPENROUTER_API_KEY or OPENROUTER_API_KEY == "SUA_CHAVE_AQUI":
        miss.append("OPENROUTER_API_KEY")
    if miss:
        return False, "Variáveis ausentes no config.py: " + ", ".join(miss)
    return True, "OK"

def validate_cnj(num: str) -> Tuple[bool, str, str]:
    d = only_digits(num)
    if len(d) != 20:
        return False, d, "Número CNJ deve conter 20 dígitos (com ou sem pontuação)."
    if STRICT_CNJ_CHECK and not cnj_checksum_ok(d):
        return False, d, "Dígito/verificação do CNJ inválido (STRICT_CNJ_CHECK=True)."
    return True, d, "OK"

def soap_consultar_processo(session: requests.Session, numero_processo: str, timeout=90,
                            movimentos=True, incluir_docs=False, debug=False) -> str:
    """
    Chama o serviço SOAP consultarProcesso e retorna o XML (texto).
    """
    envelope = f"""
    <soapenv:Envelope xmlns:soapenv="http://schemas.xmlsoap.org/soap/envelope/"
                      xmlns:ser="http://www.cnj.jus.br/servico-intercomunicacao-2.2.2/"
                      xmlns:tip="http://www.cnj.jus.br/tipos-servico-intercomunicacao-2.2.2">
        <soapenv:Header/>
        <soapenv:Body>
            <ser:consultarProcesso>
                <tip:idConsultante>{html.escape(TJ_WS_USER)}</tip:idConsultante>
                <tip:senhaConsultante>{html.escape(TJ_WS_PASS)}</tip:senhaConsultante>
                <tip:numeroProcesso>{html.escape(numero_processo)}</tip:numeroProcesso>
                <tip:movimentos>{"true" if movimentos else "false"}</tip:movimentos>
                <tip:incluirDocumentos>{"true" if incluir_docs else "false"}</tip:incluirDocumentos>
            </ser:consultarProcesso>
        </soapenv:Body>
    </soapenv:Envelope>
    """.strip()
    if debug:
        # Mascarar credenciais no log para segurança
        envelope_log = envelope.replace(TJ_WS_USER, "***USER***").replace(TJ_WS_PASS, "***PASS***")
        logger.debug("SOAP request: %s", envelope_log)
    r = session.post(TJ_WSDL_URL, data=envelope, timeout=timeout)
    r.raise_for_status()
    return r.text

def _text_of(elem: ET.Element) -> str:
    return (elem.text or "").strip() if elem is not None and elem.text else ""

def _pretty_esaj_dt(esaj_dt_str: str) -> str:
    if not esaj_dt_str:
        return esaj_dt_str
    m = re.match(r"^(\d{4})(\d{2})(\d{2})(\d{2})(\d{2})(\d{2})$", esaj_dt_str)
    if m:
        try:
            dt = datetime(*map(int, m.groups()))
            return dt.strftime("%Y-%m-%d %H:%M:%S")
        except Exception:
            pass
    return esaj_dt_str

def has_apenso_hint(xml_text: str) -> bool:
    t = xml_text.casefold()
    return (" apenso" in t) or (" apensad" in t) or (" apensamento" in t)

# ========= helpers do parser robusto (ignorando namespaces quando útil) =========
def _tagname(e: ET.Element) -> str:
    return e.tag.split('}')[-1] if isinstance(e.tag, str) else str(e)

def _iter_desc_elems(elem: ET.Element, name_endswith: str):
    """Itera descendentes cujo tag (sem ns) termina com name_endswith (case-insensitive)."""
    t = name_endswith.lower()
    for e in elem.iter():
        if _tagname(e).lower().endswith(t):
            yield e

def _all_texts(elem: ET.Element, name_endswith: str) -> List[str]:
    """Coleta todos os textos de nós cujo tag termina com name_endswith."""
    out: List[str] = []
    for e in _iter_desc_elems(elem, name_endswith):
        if e.text and e.text.strip():
            out.append(e.text.strip())
    return out

# =========================
# Parser do XML (robusto)
# =========================
def parse_xml_processo(xml_text: str) -> Dict[str, Any]:
    """
    Extrai:
      - classeProcessual (+ flag cumprimento)
      - partes AT/PA (nome + assistenciaJudiciaria)
      - decisões / despachos: codigoPaiNacional = 3 (decisão) ou 11009 (despacho)
          * aceita código em movimentoLocal e/ou movimentoLocalPai
          * coleta TODOS os <...:complemento> desse movimento
      - possivel_apenso: heurística textual
    """
    root = ET.fromstring(xml_text)
    data: Dict[str, Any] = {
        "classeProcessual": None,
        "cumprimento": False,
        "possivel_apenso": has_apenso_hint(xml_text),
        "partes": {"AT": [], "PA": []},
        "decisoes": []
    }

    # Classe processual
    dados_basicos = root.find(".//ns2:dadosBasicos", NS)
    if dados_basicos is not None:
        cls = dados_basicos.attrib.get("classeProcessual")
        data["classeProcessual"] = cls
        data["cumprimento"] = (cls in CLASSES_CUMPRIMENTO)

    # Partes por polo
    for polo_node in root.findall(".//ns2:polo", NS):
        polo = polo_node.attrib.get("polo")
        if polo not in ("AT", "PA"):
            continue
        for parte in polo_node.findall("ns2:parte", NS):
            ajg = (parte.attrib.get("assistenciaJudiciaria", "").lower() == "true")
            pessoa = parte.find("ns2:pessoa", NS)
            if pessoa is not None:
                nome = pessoa.attrib.get("nome")
                if nome:
                    data["partes"][polo].append({"nome": nome, "assistenciaJudiciaria": ajg})

    # Movimentos - AGORA COLETA TODOS OS MOVIMENTOS COM DESCRIÇÃO
    movimentos = root.findall(".//ns2:movimento", NS)
    logger.debug("Total de movimentos no XML: %d", len(movimentos))

    com_codigo = 0
    com_descricao = 0

    for mov in movimentos:
        # Códigos e descrições podem aparecer em movimentoLocal e/ou movimentoLocalPai
        cods: List[str] = []
        descrs: List[str] = []

        for ml in mov.findall("ns2:movimentoLocal", NS):
            c = ml.attrib.get("codigoPaiNacional")
            if c:
                cods.append(c)
            dsc = ml.attrib.get("descricao")
            if dsc:
                descrs.append(dsc)

            for mlp in ml.findall("ns2:movimentoLocalPai", NS):
                cp = mlp.attrib.get("codigoPaiNacional")
                if cp:
                    cods.append(cp)
                dp = mlp.attrib.get("descricao")
                if dp:
                    descrs.append(dp)

        # Fallback: qualquer nó com atributo codigoPaiNacional
        if not cods:
            for anynode in mov.iter():
                if isinstance(anynode.tag, str) and "codigoPaiNacional" in getattr(anynode, "attrib", {}):
                    cods.append(anynode.attrib.get("codigoPaiNacional"))

        if cods:
            com_codigo += 1

        # MUDANÇA: Agora coleta TODOS os movimentos que tenham descrição
        # Prioriza códigos 3 (decisão) e 11009 (despacho), mas inclui todos
        if not descrs:
            descrs = _all_texts(mov, "descricao")

        # Só adiciona se tiver descrição (para não poluir com movimentos vazios)
        if not descrs:
            continue

        com_descricao += 1

        # Pega o primeiro código disponível (se houver)
        codigo_principal = cods[0] if cods else None

        dataHora = _pretty_esaj_dt(mov.attrib.get("dataHora"))
        complementos = _all_texts(mov, "complemento")
        complemento_txt = "\n---\n".join(complementos) if complementos else ""
        descricao_final = descrs[0] if descrs else None

        data["decisoes"].append({
            "codigoPaiNacional": codigo_principal,
            "descricao": descricao_final,
            "dataHora": dataHora,
            "complemento": complemento_txt
        })

    logger.info("Movimentos com algum codigoPaiNacional: %d", com_codigo)
    logger.info("Movimentos com descrição coletados: %d", com_descricao)
    return data

# =========================
# Prompt atualizado para LLM
# =========================
def build_messages_for_llm(numero_cnj_fmt: str, dados: dict) -> list:
    resumo_json = json.dumps(dados, ensure_ascii=False, indent=2)

    sys = (
        "Você é um assistente especializado em análise processual. "
        "Produza um RELATÓRIO claro, objetivo e formal, em linguagem própria da prática forense. "
        "IMPORTANTE: Responda SEMPRE em português brasileiro, utilizando a norma culta da língua portuguesa. "

        "REGRA CRÍTICA: Todo nome de pessoa/parte deve ter **asteriscos duplos** em volta. "
        "Evite termos técnicos de programação (como true/false, AT/PA). "
        "Use expressões jurídicas completas, como 'polo ativo' e 'polo passivo'. "
        "Ao tratar de prazos, indique se o pagamento é imediato ou ao final do processo. "
        "Não escreva Tribunal de Justiça por extenso, apenas TJ-MS."
    )

    user = f"""
<contexto>
Processo: {numero_cnj_fmt}

DADOS EVIDENCIAIS (JSON):
{resumo_json}
</contexto>

<tarefas>

1. **Identificação das Partes**
   - Apresente as partes separadas por polo processual, utilizando "polo ativo" e "polo passivo".
   - OBRIGATÓRIO: Para cada parte, coloque o nome entre **asteriscos duplos** seguido de dois pontos.
   - Indique, em linguagem natural, se cada parte consta no sistema do TJ-MS como beneficiária da justiça gratuita.
   - Formato obrigatório: **Nome da Parte**: Consta no sistema como beneficiária da justiça gratuita.
   - Exemplo EXATO: **Maria da Silva**: Consta no sistema do TJ-MS como beneficiária da justiça gratuita.
   - Exemplo EXATO: **João Santos**: Não consta no sistema do TJ-MS como beneficiário da justiça gratuita.

2. **Confirmação da Gratuidade da Justiça**
   - Esclareça, para cada parte, se o sistema do TJ-MS indica a gratuidade da justiça.
   - Verifique se há decisão nos autos que conceda a gratuidade e transcreva o trecho relevante entre aspas.
   - **IMPORTANTE - IDENTIFICAÇÃO DO BENEFICIÁRIO:**
     * Analise CUIDADOSAMENTE a descrição de cada decisão/despacho para identificar QUEM é o beneficiário da justiça gratuita.
     * Procure por nomes específicos, termos como "parte autora", "requerente", "autor", "réu", "executado", etc.
     * Se a decisão mencionar nome específico de uma parte (ex: "Defiro a gratuidade a João da Silva"), associe ao nome correspondente no polo processual.
     * Se a decisão usar termo genérico mas houver apenas UMA parte naquele polo (ex: "Defiro ao autor" e só há um autor), associe àquela parte específica.
     * Se a decisão usar termo genérico e houver MÚLTIPLAS partes no polo (ex: "Defiro aos autores" e há 3 autores), considere que TODAS as partes daquele polo foram beneficiadas.
     * Se houver DÚVIDA sobre quem é o beneficiário, indique explicitamente no relatório: "⚠️ REVISÃO NECESSÁRIA: Não foi possível identificar com certeza qual parte foi beneficiada por esta decisão. Verificar manualmente."
   - Diferencie expressamente:
     (a) quando o sistema aponta gratuidade mas não há decisão confirmatória;
     (b) quando existe decisão judicial concedendo a gratuidade com beneficiário claramente identificado;
     (c) quando existe decisão judicial mas o beneficiário é ambíguo ou incerto;
     (d) quando não se identificam elementos.
   - Para cada parte, use o formato: **Nome da Parte**: [informação sobre gratuidade do sistema] + [informação sobre decisão judicial com identificação do beneficiário].
   - Exemplos de saída esperada:
     * "**João da Silva**: Consta no sistema como beneficiário. Decisão confirmatória identificou especificamente esta parte como beneficiária: 'Defiro a gratuidade ao autor João da Silva' (Despacho, 01/01/2023)."
     * "**Maria Santos**: Consta no sistema como beneficiária. Decisão deferindo gratuidade, mas ⚠️ REVISÃO NECESSÁRIA: o texto não especifica qual dos autores foi beneficiado."
     * "**Pedro Oliveira**: Não consta no sistema. Há decisão deferindo gratuidade 'aos autores', mas há 3 autores no processo. ⚠️ REVISÃO NECESSÁRIA: confirmar se esta parte específica foi beneficiada."

3. **Análise das Decisões proferidas no processo**
   - Informe apenas as decisões e despachos em que houve designação de perícia. Se não houver, informe que não há decisão ou despacho com designação de perícia.
   - Indique o valor arbitrado, quando existente (ex: R$ 500,00).
   - Esclareça quem deve arcar com o pagamento (autor, réu, Estado ou outra forma).
   - Esclareça o momento do pagamento: se a decisão determina pagamento imediato ou ao final do processo.
   - Realize a análise de conformidade com a TABELA de honorários, transcrita abaixo:
     - Considere que o juiz pode ultrapassar o limite em até 5 vezes, desde que fundamentado.
     - Classifique o valor como:
       (a) dentro da tabela;
       (b) acima da tabela, mas dentro do limite de 5 vezes com fundamentação;
       (c) acima do limite permitido;
       (d) não identificado.
     - Ao concluir, escreva obrigatoriamente:
       "Análise realizada com base na Resolução CNJ n. 232/2016, conforme redação dada pelas Resoluções n. 326/2020, n. 545/2024 e n. 599/2024."

4. **Apenso em cumprimento de sentença**
   - Se o processo não for de cumprimento de sentença, mas houve indicação de apensamento, indique isso no relatório.
   - Caso o processo seja de cumprimento de sentença e haja indícios de apensamento, finalize o relatório com a advertência:
     "Aviso: Processo de cumprimento possivelmente apensado. Recomenda-se consulta ao processo originário para confirmar a concessão da justiça gratuita."
</tarefas>

<tabela_resolucao_232>
Use a seguinte TABELA DE HONORÁRIOS como referência (valores máximos):

1. Ciências Econômicas/Contábeis
- Laudo em demanda de servidor(es) contra União/Estado/Município: R$ 300,00
- Laudo revisional envolvendo negócios bancários até 4 contratos: R$ 370,00
- Laudo revisional envolvendo negócios bancários acima de 4 contratos: R$ 630,00
- Laudo em dissolução/liquidação de sociedades civis e mercantis: R$ 830,00
- Outras: R$ 370,00

2. Engenharia/Arquitetura
- Avaliação de imóvel urbano (ABNT): R$ 430,00
- Avaliação de imóvel rural (ABNT): R$ 530,00
- Laudo estrutural/segurança de imóvel (ABNT): R$ 370,00
- Avaliação de bens fungíveis/rural/urbano (ABNT): R$ 700,00
- Ação Demarcatória: R$ 870,00
- Laudo de insalubridade/periculosidade: R$ 370,00
- Outras: R$ 370,00

3. Medicina/Odontologia
- Interdição/DNA: R$ 370,00
- Danos físicos/estéticos: R$ 370,00
- Outras: R$ 370,00

4. Psicologia: R$ 300,00
5. Serviço Social – Estudo social: R$ 300,00

6. Outras especialidades
- Avaliação comercial de bens imóveis: R$ 170,00
- Avaliação comercial por corretor: R$ 330,00
- Outras: R$ 300,00

**Regra especial (§4º do art. 2º):** O juiz pode ultrapassar o limite fixado em até 5 vezes, desde que fundamentado.
</tabela_resolucao_232>

<instruções_de_formatação>
- Estruture o relatório em seções numeradas.  
- Utilize linguagem formal, como se fosse redigido por um assistente jurídico.  
- Prefira construções como "consta no sistema", "há decisão judicial que defere", "não identificado nos autos".  
- Evite linguagem técnica de programação.  
</instruções_de_formatação>

<formato_de_saida>
A resposta deve ser redigida em **Markdown**, no formato de relatório jurídico estruturado em seções numeradas:

# Relatório - Processo XXXXXXX-XX.XXXX.X.XX.XXXX

## 1. Partes, Polos Processuais e Gratuidade da Justiça
- Apresente as partes separadas por polo processual ("polo ativo" e "polo passivo").
- Para cada parte, coloque o nome entre **asteriscos duplos** seguido de dois pontos.
- Informe, em linguagem natural:
  (a) se consta no sistema do TJ-MS como beneficiária da justiça gratuita;
  (b) se há decisão judicial confirmatória, transcrevendo o trecho relevante entre aspas E identificando SE POSSÍVEL qual parte específica foi beneficiada;
  (c) se há dúvida sobre qual parte foi beneficiada, use o marcador "⚠️ REVISÃO NECESSÁRIA";
  (d) se não há qualquer indicação.
- **REGRA CRÍTICA DE IDENTIFICAÇÃO:**
  * Se a decisão menciona nome específico, associe àquela parte.
  * Se termo genérico com UMA parte no polo, associe àquela parte.
  * Se termo genérico com MÚLTIPLAS partes no polo, considere todas beneficiadas.
  * Se AMBÍGUO ou INCERTO, marque com "⚠️ REVISÃO NECESSÁRIA".

**Formato obrigatório:**
**Nome da Parte**: [informação do sistema do TJ-MS] + [informação sobre decisão judicial com identificação clara do beneficiário].

**Exemplos de saída:**

**Polo ativo:**
- **Maria da Silva**: Consta no sistema do TJ-MS como beneficiária da justiça gratuita. Decisão confirmatória identificou especificamente esta parte: *"Defiro a gratuidade de justiça à autora Maria da Silva."* (Despacho, 01/01/2023).
- **João Santos**: Consta no sistema como beneficiário da justiça gratuita. Há decisão deferindo gratuidade aos autores de forma genérica (há 2 autores). Considera-se que ambos foram beneficiados: *"Defiro aos autores."* (Despacho, 01/01/2023).
- **Pedro Costa**: Consta no sistema como beneficiário. Há decisão deferindo gratuidade, mas ⚠️ REVISÃO NECESSÁRIA: o texto não especifica qual dos 3 autores foi beneficiado: *"Defiro ao primeiro requerente."* (Despacho, 05/01/2023).

**Polo passivo:**
- **Banco X S.A.**: Não consta no sistema nem há decisão sobre o tema.  

## 2. Análise das Decisões Proferidas no Processo
- Listar cada decisão relevante em subtópicos (por data).  
- Informar:  
  - Designação de perícia (Sim/Não).  
  - Valor arbitrado (em reais).  
  - Responsável pelo pagamento (Estado/autor/réu).  
  - Momento do pagamento (imediato/ao final do processo).  
  - Trecho da decisão entre aspas.

**Exemplo de saída:**

### Decisão de 01/01/2023
- **Designação de perícia:** Sim.  
- **Valor arbitrado:** R$ 1.500,00.  
- **Responsável pelo pagamento:** Autor.  
- **Momento do pagamento:** Ao final do processo.  
- **Trecho da decisão:** *“Defiro a produção de prova pericial, a ser custeada ao final.”*  

### Decisão de 10/02/2023
- **Designação de perícia:** Não.  
- **Valor arbitrado:** —  
- **Responsável pelo pagamento:** —  
- **Momento do pagamento:** —  
- **Trecho da decisão:** *“Indefiro o pedido de prova pericial.”*  

## 3. Processos Apensados
- Caso aplicável, incluir o aviso sobre possível apensamento. 

</formato_de_saida>

"""
    return [
        {"role": "system", "content": sys},
        {"role": "user", "content": user},
    ]

# =========================
# Cliente OpenRouter (com fallback e logs)
# =========================
def call_openrouter(messages: list, model: str = DEFAULT_MODEL, temperature=0.2, timeout=120) -> str:
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://pge-ms.lab",
        "X-Title": "Relatório TJMS",
    }
    payload = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": 20000,
    }
    r = requests.post(OPENROUTER_ENDPOINT, headers=headers, json=payload, timeout=timeout)
    logger.debug("OpenRouter status=%s", r.status_code)
    logger.debug("OpenRouter headers=%s", dict(r.headers))
    r.raise_for_status()
    j = r.json()
    # Loga só um pedaço para não poluir
    logger.debug("OpenRouter body (primeiros 600 chars): %s", json.dumps(j, ensure_ascii=False)[:600])

    # Fallbacks defensivos
    try:
        message = j["choices"][0]["message"]
        content = message.get("content", "")

        # Se tem content válido, retorna
        if content and content.strip():
            return content

        # Se não tem content, verifica se tem reasoning
        reasoning = message.get("reasoning", "")
        if reasoning and reasoning.strip():
            logger.warning("Resposta sem 'content' mas com 'reasoning'. Usando reasoning como resposta.")
            return reasoning

        # Se não tem nem content nem reasoning, tenta outros campos
        alt_content = message.get("refusal") or message.get("tool_calls") or ""
        if alt_content:
            logger.warning("Usando campo alternativo da resposta.")
            return str(alt_content)

        logger.warning("Resposta completamente vazia. Retornando mensagem de erro.")
        return "Erro: A API retornou uma resposta vazia. Tente novamente com um modelo diferente."

    except Exception:
        logger.exception("Falha ao interpretar resposta da LLM")
        return f"Erro ao processar resposta da API:\n{json.dumps(j, ensure_ascii=False, indent=2)}"

# =========================
# Pipeline alto nível
# =========================
def full_flow(numero_raw: str, model: str, diagnostic_mode=False) -> Tuple[Dict[str, Any], str]:
    ok_config, msg_config = validate_config()
    if not ok_config:
        raise RuntimeError(f"Falha na configuração: {msg_config}")
    logger.info("Configuração validada.")

    ok_cnj, d, msg_cnj = validate_cnj(numero_raw)
    if not ok_cnj:
        raise ValueError(f"CNJ inválido: {msg_cnj}")
    cnj_fmt = format_cnj(d)
    logger.info("CNJ normalizado: %s", cnj_fmt)

    session = make_session()
    xml_text = soap_consultar_processo(session, d, timeout=90, movimentos=True, incluir_docs=False, debug=(logger.level==logging.DEBUG))
    logger.info("XML recebido (%d chars).", len(xml_text))

    dados = parse_xml_processo(xml_text)
    logger.info("Dados extraídos: partes AT=%d, PA=%d; decisões=%d; classe=%s; cumprimento=%s; apenso? %s",
                len(dados["partes"]["AT"]), len(dados["partes"]["PA"]), len(dados["decisoes"]),
                dados["classeProcessual"], dados["cumprimento"], dados["possivel_apenso"])

    if diagnostic_mode:
        messages = [
            {"role": "system", "content": "Você é um analisador de sanidade. Responda sucintamente."},
            {"role": "user", "content": f"Teste: recebi JSON com AT={len(dados['partes']['AT'])}, "
                                         f"PA={len(dados['partes']['PA'])}, decs={len(dados['decisoes'])}. Diga 'OK' e ecoe os números."}
        ]
    else:
        messages = build_messages_for_llm(cnj_fmt, dados)

    rel = call_openrouter(messages, model=model)
    # reforço do aviso se for cumprimento + apenso
    if dados.get("cumprimento") and dados.get("possivel_apenso"):
        rel += "\n\nAviso: Processo de cumprimento possivelmente apensado. Talvez seja necessário consultar o processo originário para confirmar a AJG."
    logger.info("LLM respondeu com %d caracteres.", len(rel))
    return dados, rel

# =========================
# UI com log incorporado
# =========================
class UILogHandler(logging.Handler):
    def __init__(self, writer_callable):
        super().__init__()
        self.writer = writer_callable
        self.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(message)s"))

    def emit(self, record):
        try:
            self.writer(self.format(record))
        except Exception:
            pass

# ========= FUNÇÃO CORRIGIDA PARA RENDERIZAÇÃO DE MARKDOWN =========
def render_markdown_basic(text_widget: ScrolledText, markdown_text: str):
    """
    Renderiza markdown básico no widget de texto com formatação.
    Suporta: **negrito**, *itálico*, # cabeçalhos, - listas, citações com aspas.
    """
    import re

    # Limpar widget
    text_widget.delete("1.0", "end")

    margin_left = 20
    margin_right = 20
    line_spacing = 6

    # Configurar tags de formatação com justificação
    text_widget.tag_configure("h1", font=("Arial", 16, "bold"), foreground="#2E4A6B", justify="center", spacing1=0, spacing3=0)
    text_widget.tag_configure("h2", font=("Arial", 14, "bold"), foreground="#2E4A6B", justify="left", spacing1=0, spacing3=0)
    text_widget.tag_configure("h3", font=("Arial", 12, "bold"), foreground="#2E4A6B", justify="left", spacing1=0, spacing3=0)
    text_widget.tag_configure("bold", font=("Arial", 11, "bold"))
    text_widget.tag_configure("italic", font=("Arial", 10, "italic"))
    text_widget.tag_configure("quote", font=("Arial", 10, "italic"), foreground="#666666", background="#f5f5f5",
                             justify="left", lmargin1=margin_left+40, lmargin2=margin_left+40, rmargin=margin_right)
    text_widget.tag_configure("normal", font=("Arial", 10, "normal"), justify="left")
    text_widget.tag_configure("list",
        font=("Arial", 10, "normal"),
        lmargin1=margin_left+20,  # Recuo da primeira linha
        lmargin2=margin_left+20,  # Recuo das linhas subsequentes
        rmargin=margin_right,
        spacing3=line_spacing,
        justify="left")

    # Configurar espaçamento padrão apenas para texto normal (será aplicado seletivamente)
    text_widget.tag_configure("default_spacing", spacing3=line_spacing, justify="left")


    lines = markdown_text.split('\n')

    for i, line in enumerate(lines):
        # Identificar tipo de linha
        if line.startswith('### '):
            render_line(text_widget, line[4:], "h3")
        elif line.startswith('## '):
            render_line(text_widget, line[3:], "h2")
        elif line.startswith('# '):
            render_line(text_widget, line[2:], "h1")
        elif re.match(r'^\s*[-*]\s+', line):
            # Converter marcador para bullet point
            text = re.sub(r'^\s*[-*]\s+', '• ', line)
            render_line(text_widget, text, "list")
        else:
            if line.strip():
                render_line(text_widget, line, "normal")
            else:
                text_widget.insert("end", "\n")

    text_widget.see("1.0")

def render_line(text_widget: ScrolledText, line: str, base_style: str):
    """
    Renderiza uma linha com formatações inline mantidas.
    Remove marcadores markdown e aplica as formatações correspondentes.
    """
    import re

    # Estrutura para armazenar partes do texto com suas formatações
    parts = []

    # Processar a linha para identificar formatações
    remaining = line
    position = 0

    while remaining:
        # Procurar próxima formatação
        next_bold = remaining.find('**')
        next_italic = -1

        # Procurar itálico (evitando confundir com negrito)
        for i, char in enumerate(remaining):
            if char == '*' and (i == 0 or remaining[i-1] != '*') and (i+1 >= len(remaining) or remaining[i+1] != '*'):
                next_italic = i
                break

        next_quote = remaining.find('"')

        # Determinar qual vem primeiro
        positions = [(next_bold, 'bold'), (next_italic, 'italic'), (next_quote, 'quote')]
        positions = [(pos, type) for pos, type in positions if pos >= 0]

        if not positions:
            # Sem mais formatações, adicionar o resto como texto normal
            if remaining:
                parts.append((remaining, base_style))
            break

        # Ordenar por posição
        positions.sort(key=lambda x: x[0])
        next_pos, format_type = positions[0]

        # Adicionar texto antes da formatação
        if next_pos > 0:
            parts.append((remaining[:next_pos], base_style))

        # Processar a formatação
        if format_type == 'bold':
            # Encontrar o fechamento
            end = remaining.find('**', next_pos + 2)
            if end > next_pos + 2:
                content = remaining[next_pos + 2:end]
                parts.append((content, 'bold'))
                remaining = remaining[end + 2:]
            else:
                # Sem fechamento, tratar como texto normal
                parts.append((remaining[next_pos:next_pos+2], base_style))
                remaining = remaining[next_pos + 2:]

        elif format_type == 'italic':
            # Encontrar o próximo asterisco isolado
            end = next_pos + 1
            while end < len(remaining):
                if remaining[end] == '*' and (end + 1 >= len(remaining) or remaining[end + 1] != '*'):
                    content = remaining[next_pos + 1:end]
                    if content and not content.isspace():
                        parts.append((content, 'italic'))
                        remaining = remaining[end + 1:]
                    else:
                        parts.append((remaining[next_pos], base_style))
                        remaining = remaining[next_pos + 1:]
                    break
                end += 1
            else:
                # Sem fechamento
                parts.append((remaining[next_pos], base_style))
                remaining = remaining[next_pos + 1:]

        elif format_type == 'quote':
            # Encontrar o fechamento
            end = remaining.find('"', next_pos + 1)
            if end > next_pos:
                content = '"' + remaining[next_pos + 1:end] + '"'
                parts.append((content, 'quote'))
                remaining = remaining[end + 1:]
            else:
                # Sem fechamento
                parts.append((remaining[next_pos], base_style))
                remaining = remaining[next_pos + 1:]

    # Inserir as partes no widget com suas formatações
    line_start = text_widget.index("end-1c")

    for text, style in parts:
        if text:
            # Inserir com a tag diretamente
            text_widget.insert("end", text, (style,))

    text_widget.insert("end", "\n")

    # Aplicar espaçamento padrão à linha toda (Tkinter não suporta justificação)
    if base_style not in ["h1", "h2", "h3"]:
        line_end = text_widget.index("end-1c")
        text_widget.tag_add("default_spacing", line_start, line_end)

def markdown_to_rtf(markdown_text: str) -> str:
    """
    Converte markdown básico para RTF (Rich Text Format) com formatação completa.
    Suporta: **negrito**, *itálico*, # cabeçalhos, citações com aspas, listas.
    """
    import re

    # Cabeçalho RTF com configuração completa
    rtf_content = r"""{\rtf1\ansi\deff0\nouicompat\deflang1046{\fonttbl{\f0\fnil\fcharset0 Arial;}}
{\colortbl;\red0\green0\blue0;\red46\green74\blue107;\red102\green102\blue102;}
\viewkind4\uc1
\pard\sa200\sl276\slmult1\f0\fs22\lang22 """

    lines = markdown_text.split('\n')

    for line in lines:
        # Linhas vazias
        if not line.strip():
            rtf_content += r"\par "
            continue

        # Detectar e processar cabeçalhos
        if line.startswith('# '):
            # H1 - Tamanho 32 (16pt), negrito, azul, centralizado
            text = line[2:].strip()
            processed_text = process_inline_rtf(text)
            rtf_content += r"\pard\qc\sa200\sl276\slmult1{\cf2\b\fs32 " + processed_text + r"\cf0\b0\fs22\par}"
            rtf_content += r"\pard\sa200\sl276\slmult1 "

        elif line.startswith('## '):
            # H2 - Tamanho 28 (14pt), negrito, azul
            text = line[3:].strip()
            processed_text = process_inline_rtf(text)
            rtf_content += r"\pard\sa200\sl276\slmult1{\cf2\b\fs28 " + processed_text + r"\cf0\b0\fs22\par}"

        elif line.startswith('### '):
            # H3 - Tamanho 24 (12pt), negrito, azul
            text = line[4:].strip()
            processed_text = process_inline_rtf(text)
            rtf_content += r"\pard\sa200\sl276\slmult1{\cf2\b\fs24 " + processed_text + r"\cf0\b0\fs22\par}"

        # Processar listas
        elif line.strip().startswith('- ') or line.strip().startswith('* ') or line.strip().startswith('• '):
            # Remove o marcador e processa o texto
            text = re.sub(r'^\s*[-*•]\s*', '', line.strip())
            processed_text = process_inline_rtf(text)
            # Usa o bullet point unicode
            rtf_content += r"\pard\li720\fi-360\sa200\sl276\slmult1{\bullet\tab}" + processed_text + r"\par "

        else:
            # Texto normal - justificado
            processed_text = process_inline_rtf(line.strip())
            rtf_content += r"\pard\sa200\sl276\slmult1\qj " + processed_text + r"\par "

    rtf_content += r"}"
    return rtf_content

def process_inline_rtf(text: str) -> str:
    """
    Processa formatação inline (negrito, itálico, citações) para RTF.
    Mantém a estrutura e aplica formatação corretamente.
    """
    import re

    # Debug: verificar se tem marcações
    original_text = text

    # Primeiro, fazer escape de caracteres especiais RTF (exceto os marcadores markdown)
    text = text.replace('\\', '\\\\')
    text = text.replace('{', '\\{')
    text = text.replace('}', '\\}')

    # Converter caracteres acentuados para códigos RTF
    replacements = {
        'á': r"\'e1", 'à': r"\'e0", 'ã': r"\'e3", 'â': r"\'e2",
        'é': r"\'e9", 'è': r"\'e8", 'ê': r"\'ea",
        'í': r"\'ed", 'ì': r"\'ec", 'î': r"\'ee",
        'ó': r"\'f3", 'ò': r"\'f2", 'õ': r"\'f5", 'ô': r"\'f4",
        'ú': r"\'fa", 'ù': r"\'f9", 'û': r"\'fb", 'ü': r"\'fc",
        'ç': r"\'e7", 'Ç': r"\'c7",
        'Á': r"\'c1", 'À': r"\'c0", 'Ã': r"\'c3", 'Â': r"\'c2",
        'É': r"\'c9", 'È': r"\'c8", 'Ê': r"\'ca",
        'Í': r"\'cd", 'Ì': r"\'cc", 'Î': r"\'ce",
        'Ó': r"\'d3", 'Ò': r"\'d2", 'Õ': r"\'d5", 'Ô': r"\'d4",
        'Ú': r"\'da", 'Ù': r"\'d9", 'Û': r"\'db", 'Ü': r"\'dc"
    }

    for char, rtf_code in replacements.items():
        text = text.replace(char, rtf_code)

    # Processar formatações markdown
    result = ""
    i = 0

    while i < len(text):
        # Verificar negrito **texto**
        if i < len(text) - 1 and text[i:i+2] == '**':
            # Procurar o fechamento
            end = text.find('**', i + 2)
            if end != -1:
                content = text[i+2:end]
                # Aplicar negrito
                result += r"{\b " + content + r"}"
                i = end + 2
                continue

        # Verificar itálico *texto* (mas não confundir com negrito)
        if text[i] == '*' and (i == 0 or text[i-1] != '*') and (i+1 < len(text) and text[i+1] != '*'):
            # Procurar o próximo * isolado
            j = i + 1
            found = False
            while j < len(text):
                if text[j] == '*' and (j+1 >= len(text) or text[j+1] != '*'):
                    content = text[i+1:j]
                    if content and not content.isspace():
                        # Aplicar itálico
                        result += r"{\i " + content + r"}"
                        i = j + 1
                        found = True
                        break
                j += 1

            if not found:
                result += text[i]
                i += 1
            continue

        # Verificar citações "texto"
        if text[i] == '"':
            end = text.find('"', i + 1)
            if end != -1:
                content = text[i+1:end]
                # Citações em itálico com cor cinza
                result += r'{\i\cf3 "' + content + r'"\cf0\i0}'
                i = end + 1
                continue

        # Caractere normal
        result += text[i]
        i += 1

    # Debug: verificar se a conversão funcionou
    if '**' in original_text and '\\b' not in result:
        logger.warning(f"AVISO: Negrito não convertido em RTF! Original: {original_text[:50]}...")

    return result

def markdown_to_pdf(markdown_text: str, output_path: str, numero_processo: str = ""):
    """
    Converte markdown para PDF mantendo formatação (negrito, itálico, cabeçalhos).
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
        from reportlab.lib.colors import HexColor
        import re

        # Configurar documento
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )

        # Estilos personalizados
        styles = getSampleStyleSheet()

        # Estilo para título principal
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=HexColor('#2E4A6B')
        )

        # Estilo para cabeçalhos H2
        h2_style = ParagraphStyle(
            'CustomH2',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=12,
            textColor=HexColor('#2E4A6B')
        )

        # Estilo para cabeçalhos H3
        h3_style = ParagraphStyle(
            'CustomH3',
            parent=styles['Heading3'],
            fontSize=12,
            spaceBefore=16,
            spaceAfter=8,
            textColor=HexColor('#2E4A6B')
        )

        # Estilo para texto normal justificado
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6,
            alignment=TA_JUSTIFY,
            leading=15  # Espaçamento de linha 1.5
        )

        # Estilo para listas
        list_style = ParagraphStyle(
            'CustomList',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=3,
            leftIndent=20,
            bulletIndent=10,
            alignment=TA_JUSTIFY,
            leading=15
        )

        # Processar conteúdo
        story = []

        # Adicionar título com número do processo se fornecido
        if numero_processo:
            story.append(Paragraph(f"RELATÓRIO - Processo {numero_processo}", title_style))
            story.append(Spacer(1, 20))

        lines = markdown_text.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue

            # Processar diferentes tipos de linha
            if line.startswith('# '):
                text = process_inline_pdf(line[2:].strip())
                story.append(Paragraph(text, title_style))

            elif line.startswith('## '):
                text = process_inline_pdf(line[3:].strip())
                story.append(Paragraph(text, h2_style))

            elif line.startswith('### '):
                text = process_inline_pdf(line[4:].strip())
                story.append(Paragraph(text, h3_style))

            elif line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
                text = process_inline_pdf(line[2:].strip())
                story.append(Paragraph(f"• {text}", list_style))

            else:
                text = process_inline_pdf(line)
                story.append(Paragraph(text, normal_style))

        # Gerar PDF
        doc.build(story)
        return True

    except ImportError:
        return "ERROR: reportlab não instalado. Execute: pip install reportlab"
    except Exception as e:
        return f"ERROR: {str(e)}"

def process_inline_pdf(text: str) -> str:
    """
    Processa formatação inline para PDF (ReportLab).
    """
    import re

    original_text = text

    # Escapar caracteres especiais do ReportLab
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')

    # Processar negritos **texto** -> <b>texto</b>
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)

    # Processar itálicos *texto* -> <i>texto</i>
    text = re.sub(r'(?<!\*)\*([^*\s][^*]*?[^*\s])\*(?!\*)', r'<i>\1</i>', text)

    # Processar citações "texto" -> <i><font color="#666666">texto</font></i>
    text = re.sub(r'"([^"]+)"', r'<i><font color="#666666">"\1"</font></i>', text)

    # Debug: verificar se a conversão funcionou
    if '**' in original_text and '<b>' not in text:
        logger.warning(f"AVISO: Negrito não convertido em PDF! Original: {original_text[:50]}...")

    return text

# ============= FUNÇÃO DE TEMPLATE SIMPLES =============

def apply_simple_template(content: str, templates_dir: str = "templates") -> str:
    """
    Aplica template RTF simples: insere o conteúdo do relatório em um template RTF com cabeçalho/rodapé.
    MÉTODO SIMPLES: Usa o template como base e adiciona apenas texto formatado básico.
    """
    template_path = os.path.join(templates_dir, "template.rtf")

    if not os.path.exists(template_path):
        logger.info("Template não encontrado em templates/template.rtf. Usando formatação padrão.")
        return markdown_to_rtf(content)

    try:
        # Carregar template
        with open(template_path, "r", encoding="utf-8") as f:
            template_content = f.read()

        # EM VEZ DE CONVERTER PARA RTF E MISTURAR, vamos inserir texto formatado básico
        # Processar markdown para texto RTF simples SEM cabeçalhos RTF
        rtf_formatted_content = convert_markdown_to_rtf_content(content)

        # Debug: verificar se conteúdo foi gerado
        logger.debug(f"Markdown original: {len(content)} chars - {content[:200]}...")
        logger.debug(f"RTF formatado: {len(rtf_formatted_content)} chars - {rtf_formatted_content[:200]}...")

        if not rtf_formatted_content or rtf_formatted_content.strip() == r'\pard\qj Conteúdo não disponível.\par ':
            logger.warning(f"AVISO: Conteúdo RTF vazio! Markdown original: {content[:100]}...")
            # Usar fallback para conteúdo vazio
            rtf_formatted_content = r'\pard\qj ' + content.replace('\n', r'\par ') + r'\par '

        # Inserir o conteúdo no template imediatamente antes da chave de fechamento
        # Sem quebras extras para evitar página em branco
        if template_content.endswith('}'):
            result = template_content[:-1] + rtf_formatted_content + '\n}'
        else:
            result = template_content + rtf_formatted_content

        logger.info(f"Template aplicado com sucesso. Conteúdo gerado: {len(rtf_formatted_content)} chars")
        return result

    except Exception as e:
        logger.exception(f"Erro ao aplicar template: {e}")
        return markdown_to_rtf(content)  # Fallback

def convert_markdown_to_rtf_content(markdown_text: str) -> str:
    """
    Converte markdown para conteúdo RTF formatado (SEM cabeçalhos {\rtf1...)
    Retorna apenas os comandos de formatação e texto.
    """
    import re

    if not markdown_text or not markdown_text.strip():
        logger.warning("AVISO: Markdown vazio para conversão RTF!")
        return r'\pard\qj Conteúdo não disponível.\par '

    lines = markdown_text.split('\n')
    rtf_lines = []
    list_counter = 0  # Contador para listas numeradas

    for line in lines:
        line = line.strip()
        if not line:
            rtf_lines.append(r'\par ')
            list_counter = 0  # Reset contador em linha vazia
            continue

        # Cabeçalhos
        if line.startswith('### '):
            text = line[4:]
            text = process_inline_rtf_simple(text)
            rtf_lines.append(rf'\pard\qc\b\fs18 {text}\b0\fs20\par ')
            list_counter = 0
        elif line.startswith('## '):
            text = line[3:]
            text = process_inline_rtf_simple(text)
            rtf_lines.append(rf'\pard\qc\b\fs20 {text}\b0\fs20\par ')
            list_counter = 0
        elif line.startswith('# '):
            text = line[2:]
            text = process_inline_rtf_simple(text)
            rtf_lines.append(rf'\pard\qc\b\fs24 {text}\b0\fs20\par ')
            list_counter = 0

        # Listas com marcadores inteligentes
        elif line.startswith('- '):
            text = line[2:]
            text = process_inline_rtf_simple(text)
            list_counter += 1

            # Usar letras para listas pequenas, bolinhas para grandes
            if list_counter <= 26:  # a) b) c) ... z)
                letter = chr(ord('a') + list_counter - 1)
                marker = f"{letter})"
            else:  # Bolinhas para listas muito grandes
                marker = "•"

            rtf_lines.append(rf'\pard\qj\li360\fi-180 {marker} {text}\par ')

        # Citações em bloco (com recuo maior e justificação)
        elif line.startswith('> '):
            text = line[2:]
            text = process_inline_rtf_simple(text)
            rtf_lines.append(rf'\pard\qj\li720\ri720\sb120\sa120\i {text}\i0\par ')
            list_counter = 0

        # Texto normal
        else:
            text = process_inline_rtf_simple(line)
            rtf_lines.append(rf'\pard\qj {text}\par ')
            list_counter = 0

    result = '\n'.join(rtf_lines)

    # Debug: verificar se resultado não está vazio
    if not result.strip() or len(result.strip()) < 10:
        logger.warning("AVISO: Resultado RTF muito pequeno ou vazio!")
        logger.debug(f"Markdown original tinha {len(markdown_text)} chars: {markdown_text[:200]}...")
        logger.debug(f"RTF gerado: {result}")

    return result

def process_inline_rtf_simple(text: str) -> str:
    """
    Processa formatação inline para RTF de forma mais simples e segura
    """
    import re

    # Escapar caracteres especiais RTF
    text = text.replace('\\', r'\\')
    text = text.replace('{', r'\{')
    text = text.replace('}', r'\}')

    # Converter acentos
    accent_map = {
        'á': r"\'e1", 'à': r"\'e0", 'ã': r"\'e3", 'â': r"\'e2",
        'é': r"\'e9", 'è': r"\'e8", 'ê': r"\'ea",
        'í': r"\'ed", 'ì': r"\'ec", 'î': r"\'ee",
        'ó': r"\'f3", 'ò': r"\'f2", 'õ': r"\'f5", 'ô': r"\'f4",
        'ú': r"\'fa", 'ù': r"\'f9", 'û': r"\'fb",
        'ç': r"\'e7", 'Ç': r"\'c7",
        'Á': r"\'c1", 'É': r"\'c9", 'Í': r"\'cd", 'Ó': r"\'d3", 'Ú': r"\'da"
    }

    for char, code in accent_map.items():
        text = text.replace(char, code)

    # Processar negritos **texto** -> \b texto\b0
    text = re.sub(r'\*\*(.*?)\*\*', r'\\b \1\\b0 ', text)

    # Processar itálicos *texto* -> \i texto\i0
    text = re.sub(r'(?<!\*)\*([^*\s][^*]*?[^*\s])\*(?!\*)', r'\\i \1\\i0 ', text)

    # Processar citações "texto" (mantendo aspas mas em itálico)
    text = re.sub(r'"([^"]+)"', r'"\\i \1\\i0 "', text)

    return text

def rtf_to_pdf(rtf_path: str, pdf_path: str) -> bool:
    """
    Converte arquivo RTF para PDF usando LibreOffice ou método alternativo
    """
    try:
        import subprocess

        # Tentar usar LibreOffice para conversão
        try:
            # Comando para LibreOffice headless
            output_dir = os.path.dirname(pdf_path)
            cmd = [
                "soffice", "--headless", "--convert-to", "pdf",
                "--outdir", output_dir, rtf_path
            ]

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

            if result.returncode == 0:
                # LibreOffice cria arquivo com nome baseado no RTF
                rtf_name = os.path.splitext(os.path.basename(rtf_path))[0]
                generated_pdf = os.path.join(output_dir, rtf_name + ".pdf")

                if os.path.exists(generated_pdf) and generated_pdf != pdf_path:
                    # Renomear para o nome desejado
                    os.rename(generated_pdf, pdf_path)

                logger.info("RTF convertido para PDF com LibreOffice")
                return True

        except (subprocess.TimeoutExpired, FileNotFoundError) as e:
            logger.warning(f"LibreOffice não disponível ou falhou: {e}")

        # Fallback: usar reportlab para ler RTF e gerar PDF
        logger.info("Tentando conversão alternativa RTF->PDF")
        return rtf_to_pdf_reportlab(rtf_path, pdf_path)

    except Exception as e:
        logger.exception(f"Erro na conversão RTF->PDF: {e}")
        return f"Erro: {str(e)}"

def rtf_to_pdf_reportlab(rtf_path: str, pdf_path: str) -> bool:
    """
    Conversão RTF para PDF usando reportlab (método de fallback)
    DESABILITADA por estar gerando PDFs corrompidos.
    """
    logger.warning("Conversão RTF->PDF via reportlab desabilitada (gerava PDFs corrompidos)")
    return "ERRO: LibreOffice necessário para conversão RTF->PDF. Instale LibreOffice ou salve como RTF."

def markdown_to_docx(markdown_text: str, output_path: str, numero_processo: str = "") -> bool:
    """
    Converte markdown para DOCX usando python-docx com template se disponível
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE

        # Tentar carregar template personalizado
        template_path = os.path.join("templates", "template.docx")
        if os.path.exists(template_path):
            try:
                doc = Document(template_path)
                logger.info("Template DOCX carregado")

                # Limpar qualquer conteúdo existente do template para evitar página em branco
                # Manter apenas as configurações de estilo, cabeçalho e rodapé
                try:
                    # Método mais seguro: limpar parágrafos vazios e com conteúdo
                    paragraphs_to_remove = []
                    for i, paragraph in enumerate(doc.paragraphs):
                        # Marcar todos os parágrafos para remoção (vamos recriar o conteúdo)
                        paragraphs_to_remove.append(paragraph)

                    # Remover parágrafos
                    for paragraph in paragraphs_to_remove:
                        try:
                            p = paragraph._element
                            p.getparent().remove(p)
                        except:
                            pass  # Ignorar erros de remoção

                except Exception as e:
                    logger.warning(f"Erro ao limpar template: {e}")

            except Exception as e:
                logger.warning(f"Erro ao carregar template DOCX: {e}. Usando documento em branco.")
                doc = Document()
        else:
            # Criar documento em branco
            doc = Document()

        # Configurar margens se for documento novo
        if not os.path.exists(template_path):
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)

        # Processar markdown
        lines = markdown_text.split('\n')
        list_counter = 0

        # Filtrar linhas vazias consecutivas para evitar parágrafos desnecessários
        filtered_lines = []
        prev_empty = False
        for line in lines:
            line = line.strip()
            if not line:
                if not prev_empty and filtered_lines:  # Só adicionar linha vazia se já há conteúdo
                    filtered_lines.append("")
                prev_empty = True
            else:
                filtered_lines.append(line)
                prev_empty = False

        # Remover linhas vazias do início e fim
        while filtered_lines and not filtered_lines[0]:
            filtered_lines.pop(0)
        while filtered_lines and not filtered_lines[-1]:
            filtered_lines.pop()

        for i, line in enumerate(filtered_lines):
            if not line:
                # Apenas um espaço pequeno, não um parágrafo completo
                if i > 0 and i < len(filtered_lines) - 1:  # Se não é a primeira ou última linha
                    p = doc.add_paragraph()
                    p.space_after = Pt(6)  # Espaço menor
                list_counter = 0
                continue

            # Cabeçalhos (usando parágrafos com formatação manual para controle total)
            if line.startswith('### '):
                p = doc.add_paragraph()
                p.space_before = Pt(0)
                p.space_after = Pt(0)
                run = p.add_run(line[4:])
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(46, 74, 107)  # #2E4A6B
                list_counter = 0
            elif line.startswith('## '):
                p = doc.add_paragraph()
                p.space_before = Pt(0)
                p.space_after = Pt(0)
                run = p.add_run(line[3:])
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(46, 74, 107)  # #2E4A6B
                list_counter = 0
            elif line.startswith('# '):
                p = doc.add_paragraph()
                p.space_before = Pt(0)
                p.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line[2:])
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(46, 74, 107)  # #2E4A6B
                list_counter = 0

            # Listas
            elif line.startswith('- '):
                text = line[2:]
                list_counter += 1

                # Usar letras ou bolinhas
                if list_counter <= 26:
                    letter = chr(ord('a') + list_counter - 1)
                    marker = f"{letter}) "
                else:
                    marker = "• "

                p = doc.add_paragraph()
                p.add_run(marker)
                process_docx_inline_formatting(p, text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Citações
            elif line.startswith('> '):
                text = line[2:]
                p = doc.add_paragraph()
                process_docx_inline_formatting(p, text, base_italic=True)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.left_indent = Inches(0.5)
                p.right_indent = Inches(0.5)
                list_counter = 0

            # Texto normal
            else:
                p = doc.add_paragraph()
                process_docx_inline_formatting(p, line)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                list_counter = 0

        # Salvar documento
        doc.save(output_path)
        logger.info("DOCX gerado com sucesso")
        return True

    except ImportError:
        return "ERRO: python-docx não instalado. Execute: pip install python-docx"
    except Exception as e:
        logger.exception(f"Erro ao gerar DOCX: {e}")
        return f"ERRO: {str(e)}"

def markdown_to_pdf_weasyprint(markdown_text: str, output_path: str, numero_processo: str = "") -> bool:
    """
    Converte markdown para PDF usando weasyprint (alternativa ao LibreOffice)
    """
    try:
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
        import tempfile
        import os

        # Converter markdown para HTML básico
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>Relatório - {numero_processo}</title>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    font-size: 11pt;
                    line-height: 1.5;
                    margin: 1in;
                    text-align: justify;
                }}
                h1 {{
                    font-size: 16pt;
                    color: #2E4A6B;
                    text-align: center;
                    margin-bottom: 20pt;
                }}
                h2 {{
                    font-size: 14pt;
                    color: #2E4A6B;
                    margin-top: 16pt;
                    margin-bottom: 12pt;
                }}
                h3 {{
                    font-size: 12pt;
                    color: #2E4A6B;
                    margin-top: 12pt;
                    margin-bottom: 8pt;
                }}
                .quote {{
                    font-style: italic;
                    margin-left: 0.5in;
                    margin-right: 0.5in;
                    margin-top: 12pt;
                    margin-bottom: 12pt;
                    color: #666;
                }}
                .list-item {{
                    margin-bottom: 3pt;
                    margin-left: 20pt;
                }}
                @page {{
                    margin: 1in;
                    @top-center {{
                        content: "TRIBUNAL DE JUSTIÇA DE MATO GROSSO DO SUL";
                        font-size: 10pt;
                        color: #2E4A6B;
                    }}
                    @bottom-center {{
                        content: "Página " counter(page);
                        font-size: 9pt;
                    }}
                }}
            </style>
        </head>
        <body>
        """

        if numero_processo:
            html_content += f"<h1>Relatório - Processo {numero_processo}</h1>"

        # Processar markdown para HTML
        lines = markdown_text.split('\n')
        list_counter = 0

        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                html_content += "<br>"
                list_counter = 0
                continue

            # Cabeçalhos
            if line.startswith('### '):
                html_content += f"<h3>{line[4:]}</h3>"
                list_counter = 0
            elif line.startswith('## '):
                html_content += f"<h2>{line[3:]}</h2>"
                list_counter = 0
            elif line.startswith('# '):
                html_content += f"<h1>{line[2:]}</h1>"
                list_counter = 0

            # Listas
            elif line.startswith('- '):
                text = line[2:]
                list_counter += 1

                # Usar letras ou bolinhas
                if list_counter <= 26:
                    letter = chr(ord('a') + list_counter - 1)
                    marker = f"{letter}) "
                else:
                    marker = "• "

                # Processar formatação inline
                text = process_markdown_inline_html(marker + text)
                html_content += f'<div class="list-item">{text}</div>'

            # Citações
            elif line.startswith('> '):
                text = process_markdown_inline_html(line[2:])
                html_content += f'<div class="quote">{text}</div>'
                list_counter = 0

            # Texto normal
            else:
                text = process_markdown_inline_html(line)
                html_content += f"<p>{text}</p>"
                list_counter = 0

        html_content += "</body></html>"

        # Gerar PDF
        HTML(string=html_content).write_pdf(output_path)
        logger.info("PDF gerado com weasyprint")
        return True

    except ImportError:
        return "ERRO: weasyprint não instalado. Execute: pip install weasyprint"
    except Exception as e:
        logger.exception(f"Erro ao gerar PDF com weasyprint: {e}")
        return f"ERRO: {str(e)}"

def process_docx_inline_formatting(paragraph, text: str, base_italic: bool = False):
    """
    Processa formatação inline markdown e adiciona runs formatados ao parágrafo DOCX
    """
    import re

    # Limpar o parágrafo (remover runs existentes se houver)
    paragraph.clear()

    # Processar texto com regex para encontrar formatações
    parts = []
    i = 0
    current_text = ""

    while i < len(text):
        # Negrito **texto**
        if text[i:i+2] == '**':
            end_pos = text.find('**', i+2)
            if end_pos != -1:
                # Adicionar texto anterior se houver
                if current_text:
                    run = paragraph.add_run(current_text)
                    if base_italic:
                        run.italic = True
                    current_text = ""

                # Adicionar texto em negrito
                bold_text = text[i+2:end_pos]
                run = paragraph.add_run(bold_text)
                run.bold = True
                if base_italic:
                    run.italic = True

                i = end_pos + 2
                continue

        # Itálico *texto* (mas não **)
        elif text[i] == '*' and i+1 < len(text) and text[i+1] != '*':
            end_pos = text.find('*', i+1)
            if end_pos != -1 and not text[i-1:i+1] == '**' if i > 0 else True:
                # Adicionar texto anterior se houver
                if current_text:
                    run = paragraph.add_run(current_text)
                    if base_italic:
                        run.italic = True
                    current_text = ""

                # Adicionar texto em itálico
                italic_text = text[i+1:end_pos]
                run = paragraph.add_run(italic_text)
                run.italic = True

                i = end_pos + 1
                continue

        # Citações "texto"
        elif text[i] == '"':
            end_pos = text.find('"', i+1)
            if end_pos != -1:
                # Adicionar texto anterior se houver
                if current_text:
                    run = paragraph.add_run(current_text)
                    if base_italic:
                        run.italic = True
                    current_text = ""

                # Adicionar aspas e texto em itálico
                quote_text = text[i:end_pos+1]
                run = paragraph.add_run(quote_text)
                run.italic = True

                i = end_pos + 1
                continue

        # Caractere normal
        current_text += text[i]
        i += 1

    # Adicionar texto restante
    if current_text:
        run = paragraph.add_run(current_text)
        if base_italic:
            run.italic = True

def process_markdown_inline_html(text: str) -> str:
    """
    Processa formatação markdown inline para HTML
    """
    import re

    # Escapar HTML
    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    # Negrito **texto** -> <strong>texto</strong>
    text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)

    # Itálico *texto* -> <em>texto</em>
    text = re.sub(r'(?<!\*)\*([^*\s][^*]*?[^*\s])\*(?!\*)', r'<em>\1</em>', text)

    # Citações "texto"
    text = re.sub(r'"([^"]+)"', r'<em>"\1"</em>', text)

    return text

def docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """
    Converte DOCX para PDF usando LibreOffice ou outras opções
    """
    try:
        import subprocess

        # Opção 1: LibreOffice (melhor qualidade)
        try:
            output_dir = os.path.dirname(pdf_path)
            cmd = [
                "soffice", "--headless", "--convert-to", "pdf",
                "--outdir", output_dir, docx_path
            ]

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

            if result.returncode == 0:
                # LibreOffice cria arquivo com nome baseado no DOCX
                docx_name = os.path.splitext(os.path.basename(docx_path))[0]
                generated_pdf = os.path.join(output_dir, docx_name + ".pdf")

                if os.path.exists(generated_pdf) and generated_pdf != pdf_path:
                    os.rename(generated_pdf, pdf_path)

                logger.info("DOCX convertido para PDF com LibreOffice")
                return True

        except (subprocess.TimeoutExpired, FileNotFoundError) as e:
            logger.warning(f"LibreOffice não disponível: {e}")

        # Opção 2: python-docx2pdf (se disponível)
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            logger.info("DOCX convertido para PDF com docx2pdf")
            return True
        except ImportError:
            logger.warning("docx2pdf não instalado")
        except Exception as e:
            logger.warning(f"docx2pdf falhou: {e}")

        # Opção 3: Fallback - instruir usuário
        return "ERRO: Para conversão DOCX→PDF instale:\n1. LibreOffice (recomendado)\n2. pip install docx2pdf\n\nOu salve apenas como DOCX."

    except Exception as e:
        logger.exception(f"Erro na conversão DOCX→PDF: {e}")
        return f"ERRO: {str(e)}"

def extract_text_from_rtf(rtf_content: str) -> str:
    """
    Extrai texto de um arquivo RTF de forma mais inteligente
    """
    import re

    # Primeira passagem: remover definições de fontes e cores (que causam o problema)
    # Remover {\fonttbl...} completamente
    text = re.sub(r'\{\\fonttbl[^}]*\}', '', rtf_content)
    # Remover {\colortbl...} completamente
    text = re.sub(r'\{\\colortbl[^}]*\}', '', text)
    # Remover cabeçalho RTF
    text = re.sub(r'\{\\rtf1[^}]*\}', '', text)

    # Segunda passagem: simplificar comandos RTF mantendo texto
    # Remover comandos de formatação mas manter texto
    text = re.sub(r'\\pard[^\\]*', '', text)  # Parágrafo
    text = re.sub(r'\\par\s*', '\n', text)    # Quebra de linha
    text = re.sub(r'\\b\s*([^\\]+)\\b0', r'**\1**', text)  # Negrito -> markdown
    text = re.sub(r'\\i\s*([^\\]+)\\i0', r'*\1*', text)    # Itálico -> markdown
    text = re.sub(r'\\fs\d+\s*', '', text)    # Tamanho fonte
    text = re.sub(r'\\q[cjlr]\s*', '', text)  # Alinhamento
    text = re.sub(r'\\[a-z]+\d*\s*', '', text)  # Outros comandos RTF

    # Terceira passagem: limpeza final
    text = re.sub(r'[{}]', '', text)          # Chaves restantes
    text = re.sub(r'\s+', ' ', text)          # Espaços múltiplos
    text = re.sub(r'\n\s*\n', '\n\n', text)   # Quebras duplas

    return text.strip()

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Relatório TJ-MS (XML → LLM)")
        self.geometry("1080x760")

        self.var_num   = tk.StringVar()
        self.var_debug = tk.BooleanVar(value=False)

        self._dados_brutos_cache: Dict[str, Any] = {}
        self._markdown_original: str = ""  # Armazenar markdown original para exportação
        self._feedback_enviado: bool = False  # Controla se feedback já foi enviado
        self._processo_atual: str = ""  # Número do processo atual
        self._relatorio_gerado_com_sucesso: bool = False  # Controla se relatório foi gerado com sucesso

        self._build_ui()
        self._wire_logging()

        # Configurar handler para fechamento da janela
        self.protocol("WM_DELETE_WINDOW", self._on_closing)

        # Verificar chave API na inicialização (após um delay para UI carregar)
        self.after(500, self._check_api_key_on_startup)

    # --- layout ---
    def _build_ui(self):
        top = ttk.Frame(self); top.pack(fill=tk.X, padx=10, pady=8)

        ttk.Label(top, text="Número do processo (CNJ):").grid(row=0, column=0, sticky="e", padx=6, pady=4)
        self.entry_num = ttk.Entry(top, textvariable=self.var_num, width=36)
        self.entry_num.grid(row=0, column=1, sticky="w", padx=6, pady=4)

        # Registrar validação para aceitar apenas formato CNJ
        # Usar validate='focusout' em vez de 'key' para permitir colar
        vcmd = (self.register(self._validate_cnj_final), '%P')
        self.entry_num.configure(validate='focusout', validatecommand=vcmd)

        # Também adicionar validação em tempo real mais permissiva
        self.entry_num.bind('<KeyRelease>', self._on_cnj_change)
        self.entry_num.bind('<ButtonRelease>', self._on_cnj_change)

        # Adicionar dica de formato
        ttk.Label(top, text="(ex: 1234567-89.2020.1.23.4567)",
                 foreground="gray").grid(row=1, column=1, sticky="w", padx=6, pady=0)


        ttk.Checkbutton(top, text="Modo detalhado (DEBUG)", variable=self.var_debug,
                        command=self._toggle_debug).grid(row=0, column=2, sticky="w", padx=10)

        btns = ttk.Frame(self); btns.pack(fill=tk.X, padx=10, pady=6)
        ttk.Button(btns, text="Gerar Relatório", command=self._on_run).pack(side=tk.LEFT, padx=4)
        self.btn_json = ttk.Button(btns, text="Ver JSON (dados brutos)", command=self._on_view_json, state="disabled")
        self.btn_json.pack(side=tk.LEFT, padx=4)
        self.btn_save = ttk.Button(btns, text="Salvar relatório...", command=self._on_save, state="disabled")
        self.btn_save.pack(side=tk.LEFT, padx=4)
        self.btn_feedback = ttk.Button(btns, text="⚠️ Reportar Erro no Conteúdo do Relatório", command=self._on_report_error, state="disabled")
        self.btn_feedback.pack(side=tk.LEFT, padx=4)

        # Botões de configuração (lado direito)
        right_btns = ttk.Frame(btns)
        right_btns.pack(side=tk.RIGHT)

        # Botão de configuração de chave
        if KEY_MANAGER_AVAILABLE:
            self.btn_config_key = ttk.Button(right_btns, text="⚙️ Configurar Chave", command=self._on_config_key)
            self.btn_config_key.pack(side=tk.LEFT, padx=4)

        # Botão de atualização (se disponível)
        if UPDATER_AVAILABLE:
            self.btn_update = ttk.Button(right_btns, text="🔄 Verificar Atualizações", command=self._on_check_updates)
            self.btn_update.pack(side=tk.LEFT, padx=4)

        body = ttk.Panedwindow(self, orient=tk.HORIZONTAL); body.pack(fill=tk.BOTH, expand=True, padx=10, pady=8)

        left = ttk.Frame(body); body.add(left, weight=3)
        ttk.Label(left, text="Relatório:").pack(anchor="w")
        self.txt_out = ScrolledText(left, wrap="word"); self.txt_out.pack(fill=tk.BOTH, expand=True)

        right = ttk.Frame(body); body.add(right, weight=2)
        ttk.Label(right, text="Log:").pack(anchor="w")
        self.txt_log = ScrolledText(right, wrap="word"); self.txt_log.pack(fill=tk.BOTH, expand=True)

        self.status = ttk.Label(self, text="Pronto.", anchor="w")
        self.status.pack(fill=tk.X, padx=10, pady=(0,8))

    def _wire_logging(self):
        ui_handler = UILogHandler(self._append_log)
        logger.addHandler(ui_handler)

    # --- helpers UI ---
    def _append_log(self, msg: str):
        try:
            self.txt_log.insert("end", msg + "\n")
            self.txt_log.see("end")
        except tk.TclError:
            pass

    def _toggle_debug(self):
        logger.setLevel(logging.DEBUG if self.var_debug.get() else logging.INFO)
        self._append_log(f"[INFO] Nível de log ajustado para {'DEBUG' if self.var_debug.get() else 'INFO'}")

    def _set_status(self, text: str):
        self.status.configure(text=text)
        self.update_idletasks()

    def _validate_cnj_final(self, value: str) -> bool:
        """
        Validação final do CNJ (quando o usuário sai do campo)
        """
        if not value:  # Permitir campo vazio
            self.entry_num.configure(style='TEntry')
            return True

        # Limpar espaços e caracteres indesejados
        import re
        cleaned = re.sub(r'[^\d.\-]', '', value.strip())

        # Se o valor foi alterado, atualizar o campo
        if cleaned != value:
            self.var_num.set(cleaned)
            value = cleaned

        # Verificar tamanho máximo
        if len(value) > 25:
            messagebox.showwarning("CNJ Inválido",
                                 f"Número muito longo ({len(value)} caracteres). Máximo: 25 caracteres.")
            return False

        # Validar formato CNJ básico
        cnj_pattern = r'^\d{1,7}-?\d{0,2}\.?\d{0,4}\.?\d{0,1}\.?\d{0,2}\.?\d{0,4}$'

        if not re.match(cnj_pattern, value):
            messagebox.showwarning("CNJ Inválido",
                                 "Formato inválido. Use: NNNNNNN-DD.AAAA.J.TR.OOOO")
            return False

        self.entry_num.configure(style='TEntry')
        return True

    def _on_cnj_change(self, event=None):
        """
        Validação em tempo real mais permissiva (não bloqueia digitação)
        """
        value = self.var_num.get()

        if not value:
            self.entry_num.configure(style='TEntry')
            return

        # Permitir apenas números, traços e pontos (limpeza automática)
        import re
        cleaned = re.sub(r'[^\d.\-]', '', value)

        # Se foi alterado, atualizar sem disparar eventos
        if cleaned != value:
            current_pos = self.entry_num.index(tk.INSERT)
            self.var_num.set(cleaned)
            self.entry_num.icursor(min(current_pos, len(cleaned)))

        # Verificar se é muito longo
        if len(cleaned) > 25:
            # Truncar automaticamente
            truncated = cleaned[:25]
            self.var_num.set(truncated)
            self.entry_num.icursor(25)

    def _on_report_error(self):
        """Abre janela para reportar erro no relatório"""
        if not self._relatorio_gerado_com_sucesso or self._feedback_enviado:
            return

        # Criar janela de feedback
        feedback_window = tk.Toplevel(self)
        feedback_window.title("Reportar Erro no Conteúdo do Relatório")
        feedback_window.geometry("500x400")
        feedback_window.transient(self)
        feedback_window.grab_set()

        # Centralizar janela
        feedback_window.update_idletasks()
        x = (feedback_window.winfo_screenwidth() // 2) - (feedback_window.winfo_width() // 2)
        y = (feedback_window.winfo_screenheight() // 2) - (feedback_window.winfo_height() // 2)
        feedback_window.geometry(f"+{x}+{y}")

        # Conteúdo da janela
        header_frame = ttk.Frame(feedback_window)
        header_frame.pack(fill=tk.X, padx=15, pady=10)

        ttk.Label(header_frame, text="⚠️ Reportar Erro no Conteúdo do Relatório",
                 font=("Arial", 14, "bold")).pack(anchor="w")
        ttk.Label(header_frame, text=f"Processo: {self._processo_atual}",
                 font=("Arial", 10)).pack(anchor="w")

        ttk.Separator(feedback_window, orient='horizontal').pack(fill=tk.X, padx=15, pady=5)

        # Área de texto
        text_frame = ttk.Frame(feedback_window)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)

        ttk.Label(text_frame, text="Descreva o erro encontrado no relatório:").pack(anchor="w")

        error_text = ScrolledText(text_frame, height=10, wrap="word")
        error_text.pack(fill=tk.BOTH, expand=True, pady=5)

        # Placeholder que desaparece ao digitar
        placeholder_text = "Por favor, descreva o erro"
        error_text.insert("1.0", placeholder_text)
        error_text.configure(foreground="gray")

        def on_focus_in(event):
            if error_text.get("1.0", "end-1c") == placeholder_text:
                error_text.delete("1.0", "end")
                error_text.configure(foreground="black")

        def on_focus_out(event):
            if not error_text.get("1.0", "end-1c").strip():
                error_text.insert("1.0", placeholder_text)
                error_text.configure(foreground="gray")

        def on_key_press(event):
            # Se ainda tem o placeholder e uma tecla foi pressionada
            current_text = error_text.get("1.0", "end-1c")
            if current_text == placeholder_text and event.char.isprintable():
                error_text.delete("1.0", "end")
                error_text.configure(foreground="black")

        error_text.bind("<FocusIn>", on_focus_in)
        error_text.bind("<FocusOut>", on_focus_out)
        error_text.bind("<KeyPress>", on_key_press)

        # Botões
        btn_frame = ttk.Frame(feedback_window)
        btn_frame.pack(fill=tk.X, padx=15, pady=10)

        def send_error_report():
            error_description = error_text.get("1.0", "end-1c").strip()
            # Verificar se não é o placeholder e tem conteúdo suficiente
            if error_description == placeholder_text or len(error_description) < 10:
                messagebox.showwarning("Erro", "Por favor, descreva o erro com mais detalhes.")
                return

            # Enviar feedback negativo
            success = self.send_feedback_to_google_forms(
                tipo="ERRO",
                descricao=error_description,
                processo=self._processo_atual,
                modelo=DEFAULT_MODEL
            )

            if success:
                self._feedback_enviado = True
                self.btn_feedback.configure(state="disabled")
                feedback_window.destroy()
                messagebox.showinfo("Obrigado", "Erro reportado com sucesso! Obrigado pelo feedback.")
            else:
                messagebox.showerror("Erro", "Falha ao enviar feedback. Tente novamente.")

        def cancel_report():
            feedback_window.destroy()

        ttk.Button(btn_frame, text="📤 Enviar Relatório de Erro",
                  command=send_error_report).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="❌ Cancelar",
                  command=cancel_report).pack(side=tk.RIGHT, padx=5)

        # Focar na área de texto
        error_text.focus_set()
        error_text.tag_add("sel", "1.0", "end")

    def _send_automatic_positive_feedback_if_needed(self):
        """Envia feedback positivo automático se havia um relatório anterior sem feedback"""
        if self._relatorio_gerado_com_sucesso and not self._feedback_enviado:
            self.send_feedback_to_google_forms(
                tipo="SUCESSO_AUTO",
                descricao="Relatório gerado sem problemas reportados - novo relatório iniciado",
                processo=self._processo_atual,
                modelo=DEFAULT_MODEL
            )
            logger.info(f"Feedback positivo automático enviado para processo {self._processo_atual}")

            # Reset do estado apenas após enviar o feedback
            self._relatorio_gerado_com_sucesso = False
            self._feedback_enviado = False
            self._processo_atual = ""

        # Desabilitar todos os botões até o novo relatório ser gerado
        self.btn_json.configure(state="disabled")
        self.btn_save.configure(state="disabled")
        self.btn_feedback.configure(state="disabled")

    def send_feedback_to_google_forms(self, tipo: str, descricao: str, processo: str, modelo: str) -> bool:
        """Envia feedback para Google Forms - VERSÃO FUNCIONAL TESTADA"""
        import requests
        from datetime import datetime

        url = "https://docs.google.com/forms/d/e/1FAIpQLSdnbKWxgHAzaQC-RhnsSG7ojfIXz25UkaWv0xKRhMwkT0qz7A/formResponse"

        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Content-Type': 'application/x-www-form-urlencoded',
            'Origin': 'https://docs.google.com',
            'Referer': 'https://docs.google.com/forms/d/e/1FAIpQLSdnbKWxgHAzaQC-RhnsSG7ojfIXz25UkaWv0xKRhMwkT0qz7A/viewform'
        }

        # Qualquer uma dessas versões funciona (escolha uma):

        # VERSÃO 1: Com sentinel (mais segura)
        data = {
            "entry.1930801017_sentinel": "",
            "entry.1930801017": tipo,
            "entry.811378283": descricao,
            "entry.77871712": processo,
            "entry.72495185": modelo,
            "entry.864187237": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }

        # OU VERSÃO 2: Sem sentinel (mais simples)
        # data = {
        #     "entry.1930801017": tipo,
        #     "entry.811378283": descricao,
        #     "entry.77871712": processo,
        #     "entry.72495185": modelo,
        #     "entry.864187237": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        # }

        try:
            response = requests.post(url, data=data, headers=headers, timeout=30)
            return response.status_code == 200
        except Exception as e:
            print(f"Erro ao enviar feedback: {e}")
            return False

    # Use assim em qualquer lugar do seu código:
    # success = send_feedback_to_google_forms(
    #     tipo="ERRO",  # ou "SUCESSO_AUTO"
    #     descricao="Descrição do problema",
    #     processo="1234567-89.2020.1.23.4567",
    #     modelo="openai/gpt-4o-mini"
    # )

    def _on_check_updates(self):
        """Verifica se há atualizações disponíveis"""
        def check_in_thread():
            try:
                self.btn_update.configure(text="🔄 Verificando...", state="disabled")
                check_and_update(parent_window=self, silent=False)
            except Exception as e:
                logger.error(f"Erro ao verificar atualizações: {e}")
                messagebox.showerror("Erro", f"Erro ao verificar atualizações:\n{e}", parent=self)
            finally:
                try:
                    self.btn_update.configure(text="🔄 Verificar Atualizações", state="normal")
                except:
                    pass

        threading.Thread(target=check_in_thread, daemon=True).start()

    def _on_config_key(self):
        """Configura ou reconfigura a chave API"""
        if not KEY_MANAGER_AVAILABLE:
            messagebox.showerror("Erro", "Gerenciador de chaves não disponível!", parent=self)
            return

        try:
            new_key = get_api_key(parent=self, force_new=True)
            if new_key:
                # Atualiza a configuração global
                import config
                config.OPENROUTER_API_KEY = new_key

                messagebox.showinfo(
                    "Sucesso",
                    "Chave configurada com sucesso!\n\nA nova chave será usada nas próximas consultas.",
                    parent=self
                )

                # Log da mudança
                logger.info("Chave API reconfigurada pelo usuário")
            else:
                logger.info("Configuração de chave cancelada pelo usuário")
        except Exception as e:
            logger.error(f"Erro ao configurar chave: {e}")
            messagebox.showerror("Erro", f"Erro ao configurar chave:\n{e}", parent=self)

    def _check_api_key_on_startup(self):
        """Verifica se tem chave API válida na inicialização"""
        if not KEY_MANAGER_AVAILABLE:
            return True

        # Verifica se a chave atual é válida
        current_key = OPENROUTER_API_KEY
        if not current_key or current_key == "SUA_CHAVE_AQUI":
            # Pede chave ao usuário
            try:
                new_key = get_api_key(parent=self, force_new=False)
                if new_key:
                    # Atualiza configuração global
                    import config
                    config.OPENROUTER_API_KEY = new_key
                    globals()['OPENROUTER_API_KEY'] = new_key
                    logger.info("Chave API configurada na inicialização")
                    return True
                else:
                    logger.warning("Usuário cancelou configuração de chave")
                    messagebox.showwarning(
                        "Aviso",
                        "O sistema funcionará parcialmente sem a chave OpenRouter.\n\nVocê pode configurar a chave a qualquer momento no botão '⚙️ Configurar Chave'.",
                        parent=self
                    )
                    return False
            except Exception as e:
                logger.error(f"Erro na verificação de chave: {e}")
                return False

        return True

    def _on_closing(self):
        """Handler para fechamento da janela - envia feedback automático se necessário"""
        # Enviar feedback positivo automático se havia um relatório sem feedback
        if self._relatorio_gerado_com_sucesso and not self._feedback_enviado:
            self.send_feedback_to_google_forms(
                tipo="SUCESSO_AUTO",
                descricao="Relatório gerado sem problemas reportados - sistema fechado",
                processo=self._processo_atual,
                modelo=DEFAULT_MODEL
            )
            logger.info(f"Feedback positivo automático enviado ao fechar sistema para processo {self._processo_atual}")

        # Fechar a aplicação
        self.destroy()

    def _write_report(self, text: str):
        # Armazenar o markdown original para exportação
        self._markdown_original = text

        # Verificar se é um relatório de sucesso (não erro)
        is_success = not text.startswith('[ERRO') and len(text.strip()) > 50

        # Habilitar botões agora que temos conteúdo
        self.btn_json.configure(state="normal")
        self.btn_save.configure(state="normal")

        # Habilitar feedback apenas para gerações bem-sucedidas
        if is_success:
            self._relatorio_gerado_com_sucesso = True
            self._feedback_enviado = False
            self._processo_atual = self.var_num.get().strip()
            self.btn_feedback.configure(state="normal")
        else:
            self._relatorio_gerado_com_sucesso = False
            self.btn_feedback.configure(state="disabled")

        # Verifica se o texto contém markdown real (não apenas asteriscos soltos)
        import re
        has_negrito = '**' in text
        has_cabecalhos = bool(re.search(r'^#{1,3}\s', text, re.MULTILINE))
        has_italico = bool(re.search(r'\*[^*\s][^*]*[^*\s]\*', text))
        has_citacoes = '"' in text

        has_real_markdown = has_negrito or has_cabecalhos or has_italico or has_citacoes

        # Log de debug para verificar detecção
        if self.var_debug.get():
            logger.debug(f"Markdown detectado: {has_real_markdown} (negrito={has_negrito}, cabeçalhos={has_cabecalhos}, itálico={has_italico}, citações={has_citacoes})")

        if has_real_markdown and not text.startswith('[ERRO'):
            # Renderiza como markdown
            if self.var_debug.get():
                logger.debug("Renderizando como markdown")
            render_markdown_basic(self.txt_out, text)
        else:
            # Renderiza como texto simples
            if self.var_debug.get():
                logger.debug("Renderizando como texto simples")
            self.txt_out.delete("1.0", "end")
            self.txt_out.insert("end", text + "\n")
            self.txt_out.see("end")

    # --- ações ---
    def _on_run(self):
        # Enviar feedback positivo automático se havia um relatório anterior sem feedback
        self._send_automatic_positive_feedback_if_needed()

        numero = self.var_num.get().strip()
        if not numero:
            messagebox.showinfo("Atenção", "Informe um número de processo CNJ.")
            return
        self._set_status("Gerando relatório...")
        self._write_report("")

        def go():
            try:
                dados, rel = full_flow(numero, DEFAULT_MODEL, diagnostic_mode=False)
                self._dados_brutos_cache = dados
                self._write_report(rel)
                self._set_status("Concluído.")
            except Exception as e:
                logger.exception("Falha ao gerar relatório")
                self._write_report(f"[ERRO] {type(e).__name__}: {e}")
                self._set_status("Erro — ver log.")
        threading.Thread(target=go, daemon=True).start()

    def _on_view_json(self):
        if not self._dados_brutos_cache:
            messagebox.showinfo("JSON", "Não há dados carregados. Rode o teste ou gere um relatório primeiro.")
            return
        txt = json.dumps(self._dados_brutos_cache, ensure_ascii=False, indent=2)
        win = tk.Toplevel(self); win.title("JSON (dados brutos)"); win.geometry("800x600")
        box = ScrolledText(win, wrap="word"); box.pack(fill=tk.BOTH, expand=True)
        box.insert("end", txt); box.configure(state="disabled")

    def _on_save(self):
        # Usar o markdown original para manter formatação, ou texto da interface como fallback
        if self._markdown_original.strip():
            content = self._markdown_original.strip()
        else:
            content = self.txt_out.get("1.0", "end").strip()

        if not content:
            messagebox.showinfo("Salvar", "Nada para salvar.")
            return

        # Gerar nome padrão do arquivo
        numero_processo = self.var_num.get().strip()
        if numero_processo:
            # Limpar caracteres especiais do número do processo para usar no nome do arquivo
            numero_limpo = re.sub(r'[^\w\-]', '_', numero_processo)
            default_filename = f"relatório_AJG_{numero_limpo}"
        else:
            default_filename = "relatório_AJG"

        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=default_filename,
            filetypes=[
                ("Word Document", "*.docx"),
                ("PDF", "*.pdf"),
                ("Texto", "*.txt")
            ],
            title="Salvar relatório"
        )

        if path:
            try:
                if path.lower().endswith('.docx'):
                    # Gerar DOCX diretamente do markdown
                    if self.var_debug.get():
                        logger.debug(f"Salvando DOCX com {len(content)} chars")

                    result = markdown_to_docx(content, path, numero_processo)
                    if result == True:
                        messagebox.showinfo("OK", f"Relatório DOCX salvo em:\n{path}")
                    else:
                        messagebox.showerror("Erro", f"Falha ao gerar DOCX:\n{result}")

                elif path.lower().endswith('.pdf'):
                    # Gerar PDF via DOCX (método mais confiável)
                    if self.var_debug.get():
                        logger.debug(f"Salvando PDF via DOCX com {len(content)} chars")

                    # Gerar DOCX temporário
                    temp_docx_path = path.replace('.pdf', '_temp.docx')

                    try:
                        # Primeiro gerar DOCX
                        result = markdown_to_docx(content, temp_docx_path, numero_processo)
                        if result != True:
                            raise Exception(f"Falha ao gerar DOCX temporário: {result}")

                        # Converter DOCX para PDF
                        result_pdf = docx_to_pdf(temp_docx_path, path)

                        # Limpar arquivo temporário
                        if os.path.exists(temp_docx_path):
                            os.remove(temp_docx_path)

                        if result_pdf == True:
                            messagebox.showinfo("OK", f"Relatório PDF salvo em:\n{path}")
                        else:
                            messagebox.showerror("Erro", f"Falha ao gerar PDF:\n{result_pdf}")

                    except Exception as e:
                        # Limpar arquivo temporário em caso de erro
                        if os.path.exists(temp_docx_path):
                            os.remove(temp_docx_path)
                        messagebox.showerror("Erro", f"Falha ao gerar PDF via DOCX:\n{str(e)}")

                else:
                    # Salvar como texto simples
                    with open(path, "w", encoding="utf-8") as f:
                        f.write(content)
                    messagebox.showinfo("OK", f"Relatório TXT salvo em:\n{path}")

            except Exception as e:
                messagebox.showerror("Erro", f"Falha ao salvar arquivo:\n{str(e)}")
                logger.exception("Erro ao salvar relatório")

# =========================
# Entry point
# =========================
if __name__ == "__main__":
    # Verificação silenciosa de atualizações na inicialização (opcional)
    if UPDATER_AVAILABLE:
        try:
            # Verifica e aplica atualizações automaticamente em background
            def check_on_startup():
                import time
                time.sleep(2)  # Aguarda interface carregar
                check_and_update(silent=True, auto_update=True)

            threading.Thread(target=check_on_startup, daemon=True).start()
        except:
            pass

    App().mainloop()